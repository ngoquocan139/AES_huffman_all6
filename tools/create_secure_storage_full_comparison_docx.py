from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "generated_figures"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "secure_storage_full_comparison_table.docx"


ROWS = [
    ["Compression ratio", "35.015%", "29.87%", "100%", "+14.69%"],
    ["Space saving", "64.985%", "70.13%", "100%", "+7.92%"],
    ["PRD / byte match", "0.411", "0 mismatch\nbyte-exact RX", "100%", "0%"],
    ["Compression time", "~3.8641 s", "1.056 ms\nTX comp-only", "100%", "+99.97%"],
    ["Decompression time", "~0.5818 s", "0.483 ms\nRX Huffman", "100%", "+99.92%"],
    ["Encryption time", "~2.7106 s", "29.6 us\nTX AES", "100%", "+99.999%"],
    ["Decryption time", "~3.0449 s", "69.2 us\nRX AES", "100%", "+99.998%"],
    ["Compression + encryption", "~6.5747 s", "1.065 ms\nTX path", "100%", "+99.98%"],
    ["TX/RX cycles", "N/R", "53233 TX\n24222 RX", "100%", "0%"],
    ["TX/RX throughput", "N/R", "6.828 MB/s TX-in\n15.072 MB/s RX-out", "100%", "0%"],
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def format_cell(cell, bold=False, color="132033", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.bold = bold
            run.font.color.rgb = RGBColor.from_string(color)


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11)
section.page_height = Inches(8.5)
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(4)
run = title.add_run("Full Performance Comparison Table")
run.bold = True
run.font.name = "Calibri"
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x0B, 0x3B, 0x91)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(8)
subtitle_run = subtitle.add_run(
    "Secure-storage SoC compared with ECG Huffman + CBC-AES paper baseline"
)
subtitle_run.font.name = "Calibri"
subtitle_run.font.size = Pt(10.5)
subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

headers = ["Metric", "ECG CBC-AES Paper [1]", "This Work", "Correctness", "Improvement (%)"]
table = doc.add_table(rows=1, cols=len(headers))
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.style = "Table Grid"
widths = [2300, 2500, 2500, 1500, 1800]
set_table_geometry(table, widths, indent_dxa=120)
set_repeat_table_header(table.rows[0])

for idx, header in enumerate(headers):
    cell = table.rows[0].cells[idx]
    cell.text = header
    set_cell_shading(cell, "E8EEF5")
    format_cell(cell, bold=True, color="003B91", size=9.5)

for row_idx, row_data in enumerate(ROWS, start=1):
    cells = table.add_row().cells
    for col_idx, value in enumerate(row_data):
        cells[col_idx].text = value
        set_cell_margins(cells[col_idx])
        if row_idx % 2 == 0:
            set_cell_shading(cells[col_idx], "F3F7FF")
        if col_idx == 0:
            format_cell(cells[col_idx], bold=True, size=9.2)
        elif col_idx in (3, 4):
            color = "099A82" if value != "0%" else "132033"
            format_cell(cells[col_idx], bold=True, color=color, size=9.2)
        else:
            format_cell(cells[col_idx], size=9.2)

set_table_geometry(table, widths, indent_dxa=120)

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(8)
note.paragraph_format.space_after = Pt(0)
note_run = note.add_run(
    "[1] A lossless compression and encryption mechanism for remote monitoring of ECG data using Huffman coding and CBC-AES, FGCS 2019. "
    "Paper platform: MATLAB 2018a on Windows 7 64-bit, i5 2nd Gen, 8 GB RAM. N/R = not reported. "
    "RTL timing is measured from performance counters at 50 MHz on five MIT-BIH records."
)
note_run.font.name = "Calibri"
note_run.font.size = Pt(8.5)
note_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.save(OUT)
print(OUT)
